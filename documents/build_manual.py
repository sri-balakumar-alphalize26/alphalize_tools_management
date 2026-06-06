# -*- coding: utf-8 -*-
"""
Builds the NEX GENN Tool Management - User Guide.docx

Styled to match the Golden Spoon Vegetables Manual: a cover page, a contents
list, coloured "PART N" divider banners, and one section per screen using the
fixed sub-structure (What it's for / How to get there / Step by step / Good to
know) with an embedded screenshot + italic caption.

Theme uses the app's own palette: navy #2E294E + orange #F37021.

Run:  py documents/build_manual.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ----------------------------------------------------------------------------
# Paths & theme
# ----------------------------------------------------------------------------
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SHOTS = os.path.join(ROOT, "screenshots")
OUT = os.path.join(HERE, "NEX GENN Tool Management - User Guide.docx")

NAVY = RGBColor(0x2E, 0x29, 0x4E)
ORANGE = RGBColor(0xF3, 0x70, 0x21)
GREY = RGBColor(0x55, 0x55, 0x55)
BODY_FONT = "Times New Roman"
HEAD_FONT = "Times New Roman"

IMG_WIDTH_IN = 4.7  # phone screenshots are tall; keep them a sensible width

doc = Document()

# Base body style
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Recolour built-in heading styles to the app theme
for lvl, size in ((1, 18), (2, 13)):
    st = doc.styles["Heading %d" % lvl]
    st.font.name = HEAD_FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = NAVY


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def _set_shading(paragraph, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def _bottom_border(paragraph, hex_color, sz=12):
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(text="", size=11, bold=False, italic=False, color=None,
         align=None, space_before=None, space_after=None, font=BODY_FONT):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)


def part_banner(text):
    """Centered coloured part-divider banner, like Golden Spoon."""
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    _set_shading(p, "EFEDF5")          # pale navy tint
    _bottom_border(p, "F37021", sz=18)  # orange underline
    r = p.add_run(text)
    r.font.name = HEAD_FONT
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = NAVY


def add_shot(filename, caption):
    """Insert a screenshot centered, then an italic caption."""
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(IMG_WIDTH_IN))
    cap = para(caption, size=9.5, italic=True, color=GREY,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def screen(title, what, how, steps, good, shot, caption):
    """One screen section in the Golden Spoon structure."""
    h1(title)
    h2("What it's for")
    para(what)
    if how:
        h2("How to get there")
        para(how)
    if steps:
        h2("Step by step")
        bullets(steps)
    if good:
        h2("Good to know")
        bullets(good)
    add_shot(shot, caption)


# ----------------------------------------------------------------------------
# COVER
# ----------------------------------------------------------------------------
for _ in range(3):
    doc.add_paragraph()
para("NEX GENN Tool Management", size=30, bold=True, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("User Guide", size=22, bold=True, color=ORANGE,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para("Everything the app can do — from setting up the device to the very last button.",
     size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("App version 1.3.0", size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=2)
para("Powered by 369ai", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ----------------------------------------------------------------------------
# HOW TO READ + CONTENTS
# ----------------------------------------------------------------------------
h1("How to read this guide")
para("You don't have to read this front to back. If you're new, start at the "
     "beginning and work through “Getting started” — that covers connecting "
     "the device, logging in and finding your way around. After that, jump to "
     "whatever you actually need using the contents list below.")
para("Every screen in the app has its own short section here. Each one tells you "
     "what the screen is for, how to open it from the Home screen, and the exact "
     "taps to get the job done. Where it helps, there's a real screenshot so you "
     "know you're in the right place.")

h2("A couple of things before you start")
bullets([
    "Your Home screen might look shorter than the screenshots here. That's "
    "fine — your admin can hide tiles you don't need, so don't worry if "
    "something in this guide isn't on your device.",
    "Names, prices and amounts in the screenshots are sample data. Yours will "
    "show your own tools, customers and figures.",
    "Anything marked “for administrators only” needs the admin login. If "
    "that's not you, you can skip those parts.",
])

h1("What's inside")
contents = [
    ("GETTING STARTED", ["Welcome", "Connecting the device the first time",
                          "Logging in", "Picking your branch", "The Home screen"]),
    ("OPERATIONS (RENTALS)", ["New Rental", "Rental Orders"]),
    ("TOOLS, PRICING & AVAILABILITY", ["Tools & Equipment", "Pricing Rules",
                                       "Tool Categories", "Tool Availability"]),
    ("CUSTOMERS", ["Customers", "Customer ID Proofs", "Customer Ratings"]),
    ("REPORTING & MONEY", ["Sales Report", "Expenses", "Discount Details",
                           "Tax Details", "Order Reports", "Rental Dashboard"]),
    ("WRAPPING UP", ["Your Profile", "App Banners", "Logging out"]),
]
for group, items in contents:
    para(group, size=12, bold=True, color=ORANGE, space_before=8, space_after=2)
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(it)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
doc.add_page_break()

# ----------------------------------------------------------------------------
# PART 1 · GETTING STARTED
# ----------------------------------------------------------------------------
part_banner("PART 1 · GETTING STARTED")

h1("Welcome")
para("NEX GENN Tool Management is the app you use to run your tool-rental "
     "business day to day. It's a rental desk, a tool catalogue, a price list, a "
     "customer book and a set of reports rolled into one — all on a single "
     "tablet or phone.")
para("On a normal day you'll mostly live in two screens: New Rental for booking "
     "tools out to a customer, and Rental Orders for seeing what's out and "
     "bringing it back in. But there's a lot more under the surface — keeping "
     "your tools and prices up to date, watching availability, looking after "
     "customers, logging expenses and checking how the business is doing. This "
     "guide walks through all of it.")
h2("Who this is for")
bullets([
    "Counter staff — book tools out, check them back in, take payment.",
    "Managers — keep tools and prices current, watch availability, look after "
    "customers, read the reports.",
    "Administrators — set up the device, manage branches and the home-screen "
    "banners.",
])
para("The app runs on top of an Odoo back office, so anything you do here shows "
     "up there too, and the other way around. You don't need to know any of that "
     "to use it — it just means the numbers always line up.")

screen(
    "Connecting the device the first time",
    "Before anyone can log in, the tablet has to be pointed at your business's "
    "system. The Device Setup screen walks you through it from top to bottom: the "
    "Server URL (the address of your server, starting http:// or https://), the "
    "Database (once the URL is in, the list of databases loads and you pick "
    "yours), and the Admin Credentials (an admin username and password). Tap "
    "Configure Device and the tablet is linked.",
    "This is the very first screen on a brand-new device. You only ever do it "
    "once per device.",
    [
        "Type your Server URL into box 1.",
        "Wait a moment, then pick your Database from box 2 (it fills in once the "
        "URL is recognised).",
        "Enter an admin Username and Password in box 3. Tap the eye to check the "
        "password.",
        "Tap Configure Device. When it links, the app moves on to the login "
        "screen.",
    ],
    [
        "Your device has to be registered by your admin beforehand — the screen "
        "reminds you that the Device ID must be pre-registered.",
        "You only do this once. After that the app remembers the connection and "
        "goes straight to login.",
        "The first time round, your tablet may ask to use the camera — tap "
        "Allow. It's used for scanning device-registration QR codes.",
    ],
    "01_device_setup.jpg",
    "Device Setup — fill in the server, database and admin login, then tap "
    "Configure Device.",
)

screen(
    "Logging in",
    "This is the front door. Under the “Welcome back” heading there's an account "
    "card with two boxes — Username or Email and Password. Type the details your "
    "admin set up for you and the rest of the app unlocks. The app remembers you "
    "afterwards, so you won't have to do this every time — only after you log out.",
    "Open the app. If you're not already signed in, the “Welcome back” login "
    "screen is waiting for you.",
    [
        "Tap the Username or Email box and type your username — for most people "
        "that's their email address.",
        "Tap the Password box and type your password; the letters stay hidden. "
        "Tap the eye icon on the right to check what you've typed.",
        "If you want the app to remember your login on this device, turn on Auto "
        "Fill Credentials.",
        "Tap the orange Login button. If it's right, the app continues; if not, a "
        "small message pops up — check for typos and try again.",
    ],
    [
        "Auto Fill Credentials is handy on your own device — leave it off on a "
        "shared one so the next person can't sign in as you.",
        "Forgot your username or password? Your admin can look it up or reset it; "
        "there's no self-service reset.",
        "Both boxes have to be filled in or the app won't let you continue.",
    ],
    "02_login.jpg",
    "The login screen — your username is usually your email; tap the eye to "
    "check your password.",
)

screen(
    "Picking your branch",
    "If your business runs more than one branch, the app asks which one you're "
    "working from before it lets you in. Everything you then see — tools, orders, "
    "reports — belongs to that branch. The tick shows which one is currently "
    "selected.",
    "It appears straight after login (and only when you have more than one branch "
    "to choose from).",
    [
        "Read the list under “Pick a branch to continue”.",
        "Tap the branch you're working from today — it highlights and shows a "
        "tick.",
        "Tap Continue. The Home screen opens for that branch.",
    ],
    [
        "Working somewhere else tomorrow? You can switch branch again from your "
        "Profile screen.",
        "Only one branch on your account? You'll skip this screen entirely and go "
        "straight to Home.",
    ],
    "03_branch_picker.jpg",
    "Pick a branch to continue — choose where you're working from, then tap "
    "Continue.",
)

screen(
    "The Home screen",
    "Home is your launch pad. Everything in the app is one tap away from here. "
    "The tiles are sorted into sections — Operations, Tools Details & Pricing, "
    "Customers and Reporting — so you're not hunting through a long list.",
    "It's the screen you land on after logging in. The Home button at the bottom "
    "brings you back here from anywhere.",
    [
        "Across the top there's a banner that slides between any promotions your "
        "admin has set up, with today's date and the time just below it.",
        "The purple “Quick Access” strip shows which company or branch you're "
        "working in.",
        "Below that come the tiles, grouped by section. The little number next to "
        "each section heading tells you how many tiles are in it.",
        "Tap any tile to open that feature; tap the arrow at the top-left or the "
        "Home button to come back.",
        "Drag the screen down to refresh — handy if a banner or a tile has just "
        "been changed for you.",
    ],
    [
        "Three buttons sit at the bottom wherever you are: Home, Profile and "
        "Logout (Logout asks first, so a stray tap won't catch you out).",
        "Missing a tile you expected? Your admin has probably hidden it for your "
        "login — ask them to switch it back on.",
    ],
    "04_home.jpg",
    "The Home screen — tiles are grouped into sections; the bar at the bottom is "
    "always there.",
)

# ----------------------------------------------------------------------------
# PART 2 · OPERATIONS (RENTALS)
# ----------------------------------------------------------------------------
part_banner("PART 2 · OPERATIONS (RENTALS)")

screen(
    "New Rental",
    "New Rental is where a rental begins. You pick the customer, add the tools "
    "they're taking and confirm — and the order then moves along a clear track "
    "at the top: Draft → Confirmed → Checked Out → Checked In → Invoiced. The "
    "stage tells everyone exactly where the rental has got to.",
    "From the Home screen: Home → New Rental.",
    [
        "Tap New Rental on the Home screen.",
        "In the Customer box, search for the customer — or add their phone and "
        "email to create a new one on the spot. (You have to choose a customer "
        "before you can add tools.)",
        "On the Rental Lines tab, add each tool and set how many and for how long. "
        "Use the Timesheet tab for time-based hire and the Notes tab for anything "
        "the customer or staff should see.",
        "Tap Save Draft to keep it for later, or Confirm to lock the booking in "
        "and move it to the Confirmed stage.",
        "When the customer collects, the order is Checked Out; when they bring the "
        "tools back, it's Checked In; raising the bill makes it Invoiced.",
    ],
    [
        "The Responsible name shows who's handling the order — useful when more "
        "than one person works the counter.",
        "Cancel backs the whole order out if it was started by mistake.",
        "A half-finished booking saved as a Draft is waiting for you in Rental "
        "Orders, ready to pick up again.",
    ],
    "05_new_rental.jpg",
    "New Rental Order — choose the customer, add tools on Rental Lines, then "
    "Save Draft or Confirm.",
)

screen(
    "Rental Orders",
    "Rental Orders is the history of every rental. Search by order number, "
    "customer or phone, and filter by stage. Each row shows the reference "
    "(RNT/…), the customer, the dates, how many tools, the amount and a coloured "
    "status so you can see at a glance what's out, what's back and what's been "
    "paid.",
    "From the Home screen: Home → Rental Orders.",
    [
        "Tap Rental Orders on the Home screen — the newest sit at the top.",
        "Use the chips along the top — All, Draft, Confirmed, Checked Out, "
        "Checked In, Done, Invoiced, Paid, Unpaid — to show just one kind.",
        "Search by order number, customer name or phone in the box.",
        "Tap any row to open the full order, move it to the next stage, or take "
        "payment.",
        "Tap the round + at the top-right to start a new rental without going back "
        "to Home.",
    ],
    [
        "The coloured status tells you at a glance where each rental is — the "
        "same colours are used right across the app.",
        "An “Advance” line on a row means the customer has already paid part of "
        "the bill up front.",
        "Drag the list down to pull in the latest orders.",
    ],
    "06_rental_orders.jpg",
    "Rental Orders — every rental, filterable by stage, searchable, tap any one "
    "to open it.",
)

# ----------------------------------------------------------------------------
# PART 3 · TOOLS, PRICING & AVAILABILITY
# ----------------------------------------------------------------------------
part_banner("PART 3 · TOOLS, PRICING & AVAILABILITY")

screen(
    "Tools & Equipment",
    "Tools & Equipment is the full list of everything you rent out. Browse it, "
    "search it, filter by status, open any item to see its details and photo, and "
    "add new tools. Each tile shows the tool's name, its code, its rate per day "
    "and a coloured status badge.",
    "From the Home screen: Home → Tools & Equipment.",
    [
        "Tap Tools & Equipment on the Home screen.",
        "Scroll the tiles, or search by name, code, serial or barcode in the box.",
        "Tap a status chip — All, Available, Rented, Maintenance or Retired — to "
        "narrow the list.",
        "Tap any tool to see its details — photo, rate, status and description.",
        "To add a new tool, tap the round + button at the top-right and fill in "
        "its name, code, category and rate.",
    ],
    [
        "The status badge on each tile tells you instantly whether a tool is free "
        "(green Available) or out (orange Checked Out / Rented).",
        "A tool showing “No pricing” hasn't had a rate set yet — add one under "
        "Pricing Rules before you try to rent it.",
        "A clear photo makes a real difference when the counter is busy and you're "
        "picking tools by sight.",
    ],
    "07_tools.jpg",
    "Tools & Equipment — your whole catalogue, searchable and filterable by "
    "status.",
)

screen(
    "Pricing Rules",
    "Pricing Rules is the price list for everything you rent. For each product it "
    "shows the category, the number of units, the rental period (e.g. Per Day), "
    "the price and any late fee. This is where the rates on the tool tiles and on "
    "every rental come from.",
    "From the Home screen: Home → Pricing Rules.",
    [
        "Tap Pricing Rules on the Home screen.",
        "Search for a product by name, or use the Category and Period dropdowns to "
        "narrow the list.",
        "Read each row across: product, category, units, period, price and late "
        "fee.",
        "Tap a row to open it and change the rate or the late fee.",
    ],
    [
        "The Late Fee column (shown in red) is what's charged when a tool comes "
        "back later than agreed — a dash means no late fee is set.",
        "The count at the top-right (“22 of 22”) tells you how many products "
        "match your current search and filters.",
        "Change a price and it only affects rentals from then on — past orders "
        "keep the rate they were booked at.",
    ],
    "08_pricing_rules.jpg",
    "Pricing Rules — the rate and late fee for every product, searchable by name, "
    "category and period.",
)

screen(
    "Tool Categories",
    "Tool Categories is the list of groups your tools are sorted into — Hand "
    "Tools, Cleaning Equipment, Safety Equipment, Services and the like. Each row "
    "shows the category's short code and how many tools sit inside it. Those same "
    "groups become the filters you see elsewhere in the app.",
    "From the Home screen: Home → Tool Categories.",
    [
        "Tap Tool Categories on the Home screen.",
        "Read down the list — each row is a category with its code on the left "
        "and its tool count on the right.",
        "Use it to check a group exists, or to see at a glance where your tools "
        "are concentrated.",
    ],
    [
        "A category showing 0 Tools is empty — nothing has been filed under it "
        "yet.",
        "Keeping categories tidy keeps the rest of the app easy to search and "
        "filter.",
    ],
    "09_tool_categories.jpg",
    "Tool Categories — the groups your tools are sorted into, with a tool count "
    "for each.",
)

screen(
    "Tool Availability",
    "Tool Availability is the at-a-glance picture of your stock. The cards across "
    "the top total it up — total tools, how many are available, how many are "
    "checked out, total rentals and total revenue. Below, every tool is listed "
    "with its status, quantity, how many are available or out, its rentals, "
    "price, late fee and revenue.",
    "From the Home screen: Home → Tool Availability.",
    [
        "Tap Tool Availability on the Home screen.",
        "Read the summary cards for the headline numbers.",
        "Search for a tool, or use the status and category dropdowns to narrow "
        "the list.",
        "Read each row to see exactly how many of that tool are free versus out.",
        "Tap Download Excel or Download PDF to export the whole list.",
    ],
    [
        "Run this first thing to see what's free before you promise a tool to a "
        "customer.",
        "The green Avail and orange Rented badges make problem tools jump out.",
        "Availability looks after itself — it moves as tools are checked out and "
        "back in.",
    ],
    "10_tool_availability.jpg",
    "Tool Availability — totals up top, a per-tool breakdown below, exportable to "
    "Excel or PDF.",
)

# ----------------------------------------------------------------------------
# PART 4 · CUSTOMERS
# ----------------------------------------------------------------------------
part_banner("PART 4 · CUSTOMERS")

screen(
    "Customers",
    "Customers is your address book. Every person or company you rent to lives "
    "here, with their phone, email and how many rentals they've had. Search it to "
    "find someone fast, or open a record to check or update their details.",
    "From the Home screen: Home → Customers.",
    [
        "Tap Customers on the Home screen.",
        "Search by name, code, phone or email in the box at the top.",
        "Read each row — the name, contact details and the rental count on the "
        "right.",
        "Tap a customer to open their details.",
    ],
    [
        "New customers can also be added on the spot while you're starting a New "
        "Rental — they then appear here too.",
        "A customer's first initial fills the round avatar, which makes the list "
        "quick to scan.",
    ],
    "11_customers.jpg",
    "Customers — search by name, phone or email; each row shows how many rentals "
    "they've had.",
)

screen(
    "Customer ID Proofs",
    "Customer ID Proofs is where you keep a photo of each customer's ID document, "
    "front and back. It's the same customer list, opened straight onto the ID "
    "section — handy when you rent valuable tools and want proof of who took "
    "them. From the Edit Customer card you can capture or upload the ID, change "
    "it, remove it, or update the customer's other details.",
    "From the Home screen: Home → Customer ID Proofs.",
    [
        "Tap Customer ID Proofs on the Home screen.",
        "Find the customer and open their Edit Customer card.",
        "Under ID Proof – Front, tap to capture with the camera or upload a "
        "photo. Use Change to replace it or Remove to clear it.",
        "Add the back of the document under ID Proof – Back the same way.",
        "Tap Save to keep the changes.",
    ],
    [
        "The first time you use the camera, the app asks permission — tap Allow.",
        "Delete Customer at the bottom removes the whole record — use it with "
        "care.",
        "A stored ID also shows up at payment time so you can check it's the same "
        "person collecting.",
    ],
    "12_customer_id_proofs.jpg",
    "Customer ID Proofs — capture or upload the front and back of a customer's "
    "ID, then Save.",
)

screen(
    "Customer Ratings",
    "Customer Ratings lets you keep a simple note of how each customer has been "
    "to deal with — a quick reputation at a glance. Each customer carries a "
    "rating (for example Perfect), the date it was last set, and any notes. It "
    "helps the counter spot a trusted regular, or a customer to take more care "
    "with.",
    "From the Home screen: Home → Customer Ratings.",
    [
        "Tap Customer Ratings on the Home screen.",
        "Find a customer by name, code, phone or email.",
        "Read their rating badge and the date it was last set.",
        "Tap a customer to open the Customer Rating card with the full rating, "
        "the last-rated time and any notes.",
        "Tap Close when you're done.",
    ],
    [
        "Ratings are a quick judgement call — keep them honest and they'll be "
        "useful to whoever serves the customer next.",
        "Notes are the place to record why a rating was given.",
    ],
    "13_customer_ratings.jpg",
    "Customer Ratings — open a customer to see their rating, when it was set and "
    "any notes.",
)

# ----------------------------------------------------------------------------
# PART 5 · REPORTING & MONEY
# ----------------------------------------------------------------------------
part_banner("PART 5 · REPORTING & MONEY")

screen(
    "Sales Report",
    "Sales Report is the business's scoreboard. It shows total revenue and "
    "orders, the average order value, tax collected, damage charges, late fees, "
    "discounts given and the net revenue — plus your best customers and your "
    "best-earning tools. Pick a period and export the lot.",
    "From the Home screen: Home → Sales Report.",
    [
        "Tap Sales Report on the Home screen.",
        "Choose a period — Today, This Week, This Month, This Year or All Time — "
        "or set your own Custom From and To dates.",
        "Use the payment chips (All, Cash, Card, Bank, Credit) to look at just one "
        "kind of takings.",
        "Read the summary cards, then scroll to Customers Ranked and Tools Ranked.",
        "Tap Full PDF or Full Excel at the top to export the whole report, or the "
        "PDF / Excel buttons on a single table.",
    ],
    [
        "Net Revenue is what's left after discounts and costs come out — the "
        "headline figure for how the business is doing.",
        "Drag down to refresh the numbers.",
    ],
    "14_sales_report.jpg",
    "Sales Report — headline totals, best customers and best-earning tools, all "
    "in one place.",
)

screen(
    "Expenses",
    "Expenses is where you record what the business spends — fuel, repairs, "
    "supplies and the like. The cards across the top total it up for the period, "
    "and each expense moves through clear states: Draft, Submitted, Approved, "
    "Paid or Refused.",
    "From the Home screen: Home → Expenses.",
    [
        "Tap Expenses on the Home screen.",
        "Pick a period across the top — Today, Week, Month, Year or All.",
        "Use the state chips — All, Draft, Submitted, Approved, Paid, Refused — "
        "to filter.",
        "Tap + New (top-right) or Record First Expense to add one.",
        "Fill in the details and save; it then moves through the states as it's "
        "approved and paid.",
    ],
    [
        "The Total, Paid and Pending cards keep a running tally so you always know "
        "what's still owed.",
        "Keeping expenses up to date is what makes the profit figures in Sales "
        "Report trustworthy.",
    ],
    "15_expenses.jpg",
    "Expenses — totals for the period up top, with each expense tracked from "
    "Draft through to Paid.",
)

screen(
    "Discount Details",
    "Discount Details lists every order where a discount was given. The cards up "
    "top total it all — how many discounted orders, the total discount given, "
    "the average discount per order and the revenue after discount. Below, each "
    "row shows the customer, the order, who authorised it, the date, the amount "
    "knocked off and the order's status.",
    "From the Home screen: Home → Discount Details.",
    [
        "Tap Discount Details on the Home screen.",
        "Search by customer ID, order, name or who authorised it.",
        "Use the All Statuses dropdown to focus on one kind of order.",
        "Read each row to see exactly how much came off and on whose say-so.",
    ],
    [
        "The “Authorized By” column shows who approved each discount — useful "
        "for keeping giveaways in check.",
        "It's the quick way to answer “how much are we discounting, and who's "
        "giving it away?”",
    ],
    "16_discount_details.jpg",
    "Discount Details — every discounted order, who authorised it and how much "
    "came off.",
)

screen(
    "Tax Details",
    "Tax Details is the companion to Discount Details, but for tax. The cards up "
    "top total the taxed orders, the total tax collected, the average tax per "
    "order and the total revenue including tax. Each row then shows the customer, "
    "the order, the date, the subtotal, the tax added, the total and the status.",
    "From the Home screen: Home → Tax Details.",
    [
        "Tap Tax Details on the Home screen.",
        "Search by customer ID, order or name.",
        "Use the All Statuses dropdown to narrow the list.",
        "Read each row to see the subtotal, the tax added and the final total.",
    ],
    [
        "It's the one place to see exactly how much tax you've collected over a "
        "run of orders — handy at filing time.",
        "An order showing no tax (a dash, or a Draft) simply hadn't had tax "
        "applied.",
    ],
    "17_tax_details.jpg",
    "Tax Details — tax collected across your orders, with a per-order breakdown.",
)

screen(
    "Order Reports",
    "Order Reports is the master list of every order with the figures that "
    "matter. The cards up top show total orders, active rentals, total revenue "
    "and late returns. Below, each row carries the customer ID, the order, the "
    "customer, their phone, the date, the status and the amount.",
    "From the Home screen: Home → Order Reports.",
    [
        "Tap Order Reports on the Home screen.",
        "Search by customer ID, order, name or phone.",
        "Use the All Statuses dropdown to focus on one stage.",
        "Read down the list, or tap the arrow on a row to expand it.",
    ],
    [
        "Active Rentals tells you how many tools are out right now; Late Returns "
        "flags anything overdue.",
        "Where Sales Report is the summary, Order Reports is the line-by-line "
        "detail behind it.",
    ],
    "18_order_reports.jpg",
    "Order Reports — every order with its status and amount, plus headline "
    "totals up top.",
)

screen(
    "Rental Dashboard",
    "Rental Dashboard turns your rentals into a picture. The cards up top give "
    "the headline numbers — total revenue, total orders, average order, active "
    "rentals, late fees and discounts — and the chart below plots them over "
    "time. Switch what's measured and how it's grouped to see the trend you're "
    "after.",
    "From the Home screen: Home → Rental Dashboard.",
    [
        "Tap Rental Dashboard on the Home screen.",
        "Read the summary cards for the headline figures.",
        "Use the left dropdown to choose what the chart measures (e.g. Total "
        "Amount).",
        "Use the chart-type buttons to switch between bar, line and pie.",
        "Use the right dropdown to group by Month (or another period).",
    ],
    [
        "It's the quick “how are we doing?” glance, where Sales Report is the "
        "detailed dig.",
        "Tall bars are your strong periods — a fast way to spot busy and quiet "
        "months.",
    ],
    "19_rental_dashboard.jpg",
    "Rental Dashboard — headline cards and a chart of your rentals over time.",
)

# ----------------------------------------------------------------------------
# PART 6 · WRAPPING UP
# ----------------------------------------------------------------------------
part_banner("PART 6 · WRAPPING UP")

screen(
    "Your Profile",
    "Profile is your account at a glance. It shows your username and a Connected "
    "badge, your User ID, the database you're attached to, your role and your "
    "current branch. For administrators it also carries a shortcut to App "
    "Banners.",
    "Tap Profile in the bar at the bottom of the screen.",
    [
        "Tap Profile at the bottom.",
        "Check your details — User ID, Database, Role and Current Branch.",
        "Administrators can tap App Banners to manage the home-screen banners.",
    ],
    [
        "The Connected badge confirms the app is talking to your server.",
        "Working from a different branch today? This is where you switch.",
    ],
    "20_profile.jpg",
    "Your Profile — account details, your role and branch, and the App Banners "
    "shortcut for admins.",
)

h1("App Banners")
para("App Banners (for administrators only) controls the sliding pictures that "
     "appear across the top of the Home screen — the promotions everyone sees "
     "when they open the app. You reach it from the Profile screen.")
h2("Good to know")
bullets([
    "Banners are a great spot for a seasonal offer or a reminder to staff.",
    "Changes show up for everyone the next time they open or refresh Home.",
    "If you don't see App Banners on your Profile, you're not signed in as an "
    "administrator.",
])

screen(
    "Logging out",
    "Logging out signs you off the device so the next person has to use their own "
    "login. Because the app remembers you between sessions, logging out is the "
    "deliberate way to hand the tablet over.",
    "Tap Logout in the bar at the bottom of the screen.",
    [
        "Tap Logout at the bottom.",
        "A small box asks “Are you sure you want to logout?”",
        "Tap OK to sign out, or Cancel to stay.",
        "After logging out you're taken back to the login screen.",
    ],
    [
        "The confirmation step means a stray tap won't sign you out by accident.",
        "On a shared device, always log out at the end of your shift so your "
        "actions stay tied to you.",
    ],
    "21_logout.jpg",
    "Logging out — the app asks first, so a stray tap won't sign you out.",
)

# ----------------------------------------------------------------------------
doc.save(OUT)
print("Saved:", OUT)
print("Paragraphs:", len(doc.paragraphs))
img_count = len(doc.inline_shapes)
print("Embedded images:", img_count)
