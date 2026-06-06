# -*- coding: utf-8 -*-
"""
Frames and resizes every screenshot in the EXISTING (hand-edited) manual to
match the Golden Spoon manual, without touching any other content:

  * each screenshot resized to 4.0" wide (aspect preserved)
  * each wrapped in a single-cell table with a thin black border (sz=4, auto)
  * table is 100% content width, centered; image centered inside
  * the italic caption stays in the next paragraph, outside the frame

Edits the file IN PLACE. A safety backup copy is made first.

Run:  py documents/frame_images.py
"""

import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "NEX GENN Tool Management - User Guide.docx")
BACKUP = os.path.join(HERE, "NEX GENN Tool Management - User Guide (backup before framing).docx")

IMG_WIDTH = Inches(4.0)


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def _build_empty_frame():
    """Build a single-cell bordered table with an empty cell; return (tbl, tc)."""
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblPr.append(_el("w:tblW", **{"w:w": "5000", "w:type": "pct"}))
    tblPr.append(_el("w:jc", **{"w:val": "center"}))
    borders = OxmlElement("w:tblBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        borders.append(_el(side, **{"w:val": "single", "w:sz": "4",
                                     "w:space": "0", "w:color": "auto"}))
    tblPr.append(borders)
    cellmar = OxmlElement("w:tblCellMar")
    cellmar.append(_el("w:left", **{"w:w": "10", "w:type": "dxa"}))
    cellmar.append(_el("w:right", **{"w:w": "10", "w:type": "dxa"}))
    tblPr.append(cellmar)
    tbl.append(tblPr)

    grid = OxmlElement("w:tblGrid")
    grid.append(_el("w:gridCol", **{"w:w": "9000"}))
    tbl.append(grid)

    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcPr.append(_el("w:tcW", **{"w:w": "0", "w:type": "auto"}))
    tc.append(tcPr)
    tr.append(tc)
    tbl.append(tr)
    return tbl, tc


def _center_paragraph(p_element):
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")


def main():
    if not os.path.exists(DOCX):
        raise SystemExit("Manual not found: %s" % DOCX)

    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    imgs_before = len(doc.inline_shapes)
    paras_before = len(doc.paragraphs)

    # 1) resize each screenshot to 4.0" wide, aspect preserved
    for shape in doc.inline_shapes:
        w, h = shape.width, shape.height
        if w and h:
            ratio = h / w
            shape.width = IMG_WIDTH
            shape.height = Inches(4.0 * ratio)

    # 2) wrap each image paragraph in a bordered single-cell table.
    body = doc.element.body
    DRAW = qn("w:drawing")
    image_ps = [p for p in body.findall(qn("w:p")) if p.findall(".//" + DRAW)]

    wrapped = 0
    for p in image_ps:
        tbl, tc = _build_empty_frame()
        p.addprevious(tbl)      # put the (empty) table where the image paragraph is
        _center_paragraph(p)    # center the image inside the cell
        tc.append(p)            # move the image paragraph into the cell
        wrapped += 1

    doc.save(DOCX)

    # verify
    check = Document(DOCX)
    tbls = check.element.body.findall(qn("w:tbl"))
    tbls_with_img = sum(1 for t in tbls if t.findall(".//" + DRAW))
    print("Backup saved:", BACKUP)
    print("Images resized & framed:", wrapped)
    print("Inline images: %d (was %d)" % (len(check.inline_shapes), imgs_before))
    print("Tables in body: %d  (with an image: %d)" % (len(tbls), tbls_with_img))
    print("Body paragraphs: %d (was %d, image paras moved into cells)" %
          (len(check.paragraphs), paras_before))


if __name__ == "__main__":
    main()
