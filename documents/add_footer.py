# -*- coding: utf-8 -*-
"""
Adds a page footer to the EXISTING (already hand-edited) manual without
touching its content:

    NEX GENN Tool Management — User Guide    |    Page {PAGE} of {NUMPAGES}

Centered, 10pt, grey #888888, Times New Roman — matching the Golden Spoon footer.
The PAGE / NUMPAGES are live Word fields, so "Page X of Y" updates automatically.

Run:  py documents/add_footer.py
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "NEX GENN Tool Management - User Guide.docx")
BACKUP = os.path.join(HERE, "NEX GENN Tool Management - User Guide (backup before footer).docx")

FOOTER_TEXT = "NEX GENN Tool Management — User Guide    |    Page "
GREY = RGBColor(0x88, 0x88, 0x88)
FONT = "Times New Roman"
SIZE = Pt(10)


def _style_run(run):
    run.font.name = FONT
    run.font.size = SIZE
    run.font.color.rgb = GREY
    # ensure the font name sticks for all script ranges
    rPr = run._element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)


def _add_field(paragraph, field):
    """Append a live Word field (e.g. PAGE / NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    _style_run(run)
    r = run._element

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " %s " % field
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, sep, cached, end):
        r.append(el)


def build_footer(footer):
    footer.is_linked_to_previous = False
    # reuse the first existing paragraph, clearing it
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(FOOTER_TEXT)
    _style_run(r)
    _add_field(p, "PAGE")
    r = p.add_run(" of ")
    _style_run(r)
    _add_field(p, "NUMPAGES")


def main():
    if not os.path.exists(DOCX):
        raise SystemExit("Manual not found: %s" % DOCX)

    # Safety: keep a copy of the hand-edited file before re-saving.
    shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    imgs_before = len(doc.inline_shapes)
    paras_before = len(doc.paragraphs)

    for section in doc.sections:
        build_footer(section.footer)

    doc.save(DOCX)

    # Verify nothing in the body was lost.
    check = Document(DOCX)
    print("Backup saved:", BACKUP)
    print("Footer added to", len(doc.sections), "section(s).")
    print("Body paragraphs: %d (was %d)" % (len(check.paragraphs), paras_before))
    print("Images: %d (was %d)" % (len(check.inline_shapes), imgs_before))


if __name__ == "__main__":
    main()
